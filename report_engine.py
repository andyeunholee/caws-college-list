# -*- coding: utf-8 -*-
"""
report_engine.py
================
Rendering engine that reproduces — pixel-for-pixel in Word — the exact visual
format of the Elite Prep "Comprehensive College Admissions Strategy Report"
template (reference file:
    Eng-Sydney Bang-La_College_List_Strategy_CE_06.30.2026.docx).

The engine is fully data-driven: feed it a `student` dict and a `content` dict
(see generate_college_report.py) and it emits an identically-formatted .docx.

Requires: python-docx  (pip install python-docx)
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips


# ---------------------------------------------------------------------------
# Palette / constants (all taken verbatim from the reference document)
# ---------------------------------------------------------------------------
FONT      = "Malgun Gothic"
NAVY      = "1F3864"
RED       = "7D2230"
GOLD      = "B7791F"
GREEN     = "2F7D32"
TEXT      = "222222"
GRAY      = "666666"
WHITE     = "FFFFFF"

CONTENT_W = 9360           # usable width in dxa (8.5in page - 1in margins each side)

# Per-tier styling: accent bar, header/label background, alternating-row tint,
# probability-text colour, and the light box background.
TIER = {
    "reach":  dict(accent=RED,   head=RED,   alt="F3E7E9", prob=RED,   box="F3E7E9"),
    "match":  dict(accent=GOLD,  head=GOLD,  alt="FBF1DD", prob=GOLD,  box="FBF1DD"),
    "safety": dict(accent=GREEN, head=GREEN, alt="E7F1E8", prob=GREEN, box="E7F1E8"),
    "navy":   dict(accent=NAVY,  head=NAVY,  alt="F4F4F4", prob=NAVY,  box="E8ECF4"),
}


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------
def _sub(parent, tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn("w:" + k), str(v))
    parent.append(el)
    return el


def _run(paragraph, text, *, bold=False, color=TEXT, size=18, italic=False):
    """Append a run with explicit formatting. `size` is in half-points (18 = 9pt)."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    _sub(rPr, "w:rFonts", ascii=FONT, cs=FONT, eastAsia=FONT, hAnsi=FONT)
    if bold:
        _sub(rPr, "w:b"); _sub(rPr, "w:bCs")
    else:
        _sub(rPr, "w:b", val="false"); _sub(rPr, "w:bCs", val="false")
    if italic:
        _sub(rPr, "w:i"); _sub(rPr, "w:iCs")
    _sub(rPr, "w:color", val=color)
    _sub(rPr, "w:sz", val=size); _sub(rPr, "w:szCs", val=size)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    paragraph._p.append(r)
    return r


def _para(container, *, after=None, before=None, line=None, align=None,
          shd=None, bottom_border=None, style=None):
    """Create a bare paragraph inside `container` (a Document or a _Cell)."""
    p = container.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    if style:
        _sub(pPr, "w:pStyle", val=style)
    if bottom_border:
        pBdr = _sub(pPr, "w:pBdr")
        _sub(pBdr, "w:bottom", val="single", color=bottom_border[0],
             sz=bottom_border[1], space=bottom_border[2])
    if shd:
        _sub(pPr, "w:shd", fill=shd, color="auto", val="clear")
    spacing = {}
    if after is not None:  spacing["after"] = after
    if before is not None: spacing["before"] = before
    if line is not None:   spacing["line"] = line
    if spacing:
        _sub(pPr, "w:spacing", **spacing)
    if align:
        _sub(pPr, "w:jc", val=align)
    return p


def _tbl_borders(tbl, *, outline="D5D5D5", inside="E5E5E5",
                 left_accent=None, inside_v_white=False, none_all=False):
    tblPr = tbl._tbl.tblPr
    b = _sub(tblPr, "w:tblBorders")
    def edge(tag, val, color, sz):
        _sub(b, "w:" + tag, val=val, color=color, sz=sz)
    if none_all:
        for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge(e, "none", WHITE, 0)
        if left_accent:
            b.remove(b.find(qn("w:left")))
            edge("left", "single", left_accent[0], left_accent[1])
        if inside_v_white:
            b.remove(b.find(qn("w:insideV")))
            edge("insideV", "single", WHITE, 4)
        return
    for e in ("top", "left", "bottom", "right"):
        edge(e, "single", outline, 4)
    edge("insideH", "single", inside, 4)
    edge("insideV", "single", inside, 4)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:shd"))
    if old is not None:
        tcPr.remove(old)
    _sub(tcPr, "w:shd", fill=fill, color="auto", val="clear")


def _cell_margins(cell, top=60, left=90, bottom=60, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    m = _sub(tcPr, "w:tcMar")
    _sub(m, "w:top", type="dxa", w=top)
    _sub(m, "w:left", type="dxa", w=left)
    _sub(m, "w:bottom", type="dxa", w=bottom)
    _sub(m, "w:right", type="dxa", w=right)


def _vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    _sub(tcPr, "w:vAlign", val="center")


def _cell_left_border(cell, color, sz):
    tcPr = cell._tc.get_or_add_tcPr()
    b = _sub(tcPr, "w:tcBorders")
    _sub(b, "w:top", val="none", color=WHITE, sz=0)
    _sub(b, "w:left", val="single", color=color, sz=sz)
    _sub(b, "w:bottom", val="none", color=WHITE, sz=0)
    _sub(b, "w:right", val="none", color=WHITE, sz=0)


def _clear_cell(cell):
    """Remove the default empty paragraph python-docx puts in a fresh cell."""
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def _set_grid(tbl, widths):
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc in list(grid):
        grid.remove(gc)
    for w in widths:
        _sub(grid, "w:gridCol", w=w)
    tblPr = tbl._tbl.tblPr
    _sub(tblPr, "w:tblW", type="dxa", w=sum(widths))


def _keep_header(row):
    trPr = row._tr.get_or_add_trPr()
    _sub(trPr, "w:tblHeader")


def _no_split(row):
    trPr = row._tr.get_or_add_trPr()
    _sub(trPr, "w:cantSplit")


# ===========================================================================
#  ReportBuilder
# ===========================================================================
class ReportBuilder:
    def __init__(self):
        self.doc = Document()
        self._page_setup()

    # -- document scaffolding ------------------------------------------------
    def _page_setup(self):
        sec = self.doc.sections[0]
        sec.page_width  = Twips(12240)
        sec.page_height = Twips(15840)
        sec.top_margin = sec.bottom_margin = Twips(1440)
        sec.left_margin = sec.right_margin = Twips(1440)
        sec.header_distance = Twips(708)
        sec.footer_distance = Twips(708)

    def build_header(self, right_text):
        hdr = self.doc.sections[0].header
        hdr.is_linked_to_previous = False
        p = hdr.paragraphs[0]
        pPr = p._p.get_or_add_pPr()
        pBdr = _sub(pPr, "w:pBdr")
        _sub(pBdr, "w:bottom", val="single", color="D5D5D5", sz=6, space=4)
        tabs = _sub(pPr, "w:tabs")
        _sub(tabs, "w:tab", val="right", pos=CONTENT_W)
        _sub(pPr, "w:spacing", after=0)
        _run(p, "ELITE PREP", bold=True, color=NAVY, size=16)
        _run(p, "\t" + right_text, bold=False, color=GRAY, size=16)

    def build_footer(self, text):
        ftr = self.doc.sections[0].footer
        ftr.is_linked_to_previous = False
        p = ftr.paragraphs[0]
        pPr = p._p.get_or_add_pPr()
        pBdr = _sub(pPr, "w:pBdr")
        _sub(pBdr, "w:top", val="single", color="D5D5D5", sz=6, space=4)
        _sub(pPr, "w:spacing", before=60)
        _sub(pPr, "w:jc", val="center")
        _run(p, text, bold=False, color=GRAY, size=15)
        # PAGE field
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        _sub(rPr, "w:rFonts", ascii=FONT, cs=FONT, eastAsia=FONT, hAnsi=FONT)
        _sub(rPr, "w:color", val=GRAY)
        _sub(rPr, "w:sz", val=15); _sub(rPr, "w:szCs", val=15)
        r.append(rPr)
        for t, attr in (("begin", None), (None, "PAGE"), ("separate", None), ("end", None)):
            if attr:
                it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
                it.text = attr; r.append(it)
            else:
                fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), t); r.append(fc)
        p._p.append(r)

    # -- primitives ----------------------------------------------------------
    def spacer(self, after=120):
        p = _para(self.doc, after=after)
        _run(p, "\n", color=TEXT, size=20)

    def page_break(self):
        p = self.doc.add_paragraph()
        r = OxmlElement("w:r")
        br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
        r.append(br); p._p.append(r)

    def heading(self, text):
        p = _para(self.doc, style="Heading1", after=140, before=260,
                  bottom_border=(NAVY, 18, 4))
        _run(p, text, bold=True, color=NAVY, size=30)

    def body(self, runs, after=120, line=276):
        """runs: list of (text, bold, color, size)."""
        p = _para(self.doc, after=after, before=0, line=line, align="left")
        for text, bold, color, size in runs:
            _run(p, text, bold=bold, color=color, size=size)
        return p

    def _new_table(self, rows, widths):
        tbl = self.doc.add_table(rows=rows, cols=len(widths))
        _set_grid(tbl, widths)
        return tbl

    # -- boxes ---------------------------------------------------------------
    def box(self, paragraphs, *, bg, accent):
        """A single-cell callout box with a coloured left accent bar.
        paragraphs: list of (runs, spacing_after) where runs is list of
        (text, bold, color, size)."""
        tbl = self._new_table(1, [CONTENT_W])
        _tbl_borders(tbl, none_all=True, left_accent=(accent, 28))
        row = tbl.rows[0]
        _no_split(row)
        cell = row.cells[0]
        _clear_cell(cell)
        _shade(cell, bg)
        _cell_margins(cell, top=120, left=200, bottom=120, right=160)
        for runs, sp in paragraphs:
            p = cell.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            _sub(pPr, "w:spacing", after=sp, before=0, line=276)
            _sub(pPr, "w:jc", val="left")
            for text, bold, color, size in runs:
                _run(p, text, bold=bold, color=color, size=size)
        return tbl

    def tier_badges(self, labels=("Reach  ·  Est. ≤ 20%", "Match  ·  Est. 21–55%", "Safety  ·  Est. ≥ 60%")):
        tbl = self._new_table(1, [3120, 3120, 3120])
        _tbl_borders(tbl, none_all=True, inside_v_white=True)
        specs = [(labels[0], RED, "F3E7E9"),
                 (labels[1], GOLD, "FBF1DD"),
                 (labels[2], GREEN, "E7F1E8")]
        for cell, (label, color, bg) in zip(tbl.rows[0].cells, specs):
            _clear_cell(cell)
            _cell_left_border(cell, color, 24)
            _shade(cell, bg)
            _cell_margins(cell, top=80, left=160, bottom=80, right=120)
            p = cell.add_paragraph()
            _sub(p._p.get_or_add_pPr(), "w:spacing", after=0, line=230)
            _run(p, label, bold=True, color=color, size=17)

    def tier_label(self, tier, text):
        p = _para(self.doc, after=140, before=220, line=300, shd=TIER[tier]["head"])
        _run(p, "  " + text, bold=True, color=WHITE, size=24)

    # -- the college list tables (No. / College / State / Est.) --------------
    def list_table(self, tier, rows, headers=("No.", "College", "State", "Est. Admit Probability")):
        """rows: list of (no, college, state, prob)."""
        t = TIER[tier]
        widths = [560, 5600, 800, 2400]
        tbl = self._new_table(1 + len(rows), widths)
        _tbl_borders(tbl)
        # header
        hdr = tbl.rows[0]
        _keep_header(hdr)
        _jc = ("center", "left", "center", "center")
        for cell, (label, jc) in zip(hdr.cells, zip(headers, _jc)):
            _clear_cell(cell); _shade(cell, t["head"]); _vcenter(cell)
            _cell_margins(cell, top=60, left=90, bottom=60, right=90)
            p = cell.add_paragraph()
            _sub(p._p.get_or_add_pPr(), "w:spacing", after=0, line=240)
            _sub(p._p.get_or_add_pPr(), "w:jc", val=jc)
            _run(p, label, bold=True, color=WHITE, size=18)
        # data rows
        for i, (no, college, state, prob) in enumerate(rows):
            fill = WHITE if i % 2 == 0 else t["alt"]
            cells = tbl.rows[i + 1].cells
            data = [(str(no), "center", False, TEXT),
                    (college, "left", False, TEXT),
                    (state, "center", False, TEXT),
                    (prob, "center", True, t["prob"])]
            for cell, (val, jc, bold, color) in zip(cells, data):
                _clear_cell(cell); _shade(cell, fill); _vcenter(cell)
                _cell_margins(cell, top=46, left=90, bottom=46, right=90)
                p = cell.add_paragraph()
                _sub(p._p.get_or_add_pPr(), "w:spacing", after=0, line=230)
                _sub(p._p.get_or_add_pPr(), "w:jc", val=jc)
                _run(p, val, bold=bold, color=color, size=18)
        return tbl

    # -- generic navy-header data table (info / ED / EA / abbrev) ------------
    def _generic_table(self, widths, header, rows, *, header_bg=NAVY):
        """header: list of (label, jc). rows: list of list of
        (text, jc, bold, color). Alternating F4F4F4 rows."""
        tbl = self._new_table(1 + len(rows), widths)
        _tbl_borders(tbl)
        hdr = tbl.rows[0]
        _keep_header(hdr)
        for cell, (label, jc) in zip(hdr.cells, header):
            _clear_cell(cell); _shade(cell, header_bg); _vcenter(cell)
            _cell_margins(cell, top=60, left=80, bottom=60, right=80)
            p = cell.add_paragraph()
            _sub(p._p.get_or_add_pPr(), "w:spacing", after=0, line=230)
            _sub(p._p.get_or_add_pPr(), "w:jc", val=jc)
            _run(p, label, bold=True, color=WHITE, size=18)
        for i, row in enumerate(rows):
            fill = WHITE if i % 2 == 0 else "F4F4F4"
            cells = tbl.rows[i + 1].cells
            for cell, (text, jc, bold, color) in zip(cells, row):
                _clear_cell(cell); _shade(cell, fill); _vcenter(cell)
                _cell_margins(cell, top=60, left=80, bottom=60, right=80)
                p = cell.add_paragraph()
                _sub(p._p.get_or_add_pPr(), "w:spacing", after=0, line=230)
                _sub(p._p.get_or_add_pPr(), "w:jc", val=jc)
                _run(p, text, bold=bold, color=color, size=18)
        return tbl

    # -- title block ---------------------------------------------------------
    def title_block(self, student_name, cycle_line,
                    subtitle="College Admissions Strategy Report",
                    title="Comprehensive College Admissions Strategy Report"):
        # 1. ELITE PREP (letter-spaced)
        p = _para(self.doc, after=0, before=600, align="left")
        r = _run(p, "ELITE PREP", bold=True, color=NAVY, size=30)
        _sub(r.find(qn("w:rPr")), "w:spacing", val=60)
        # 2. sub-title with gold underline
        p = _para(self.doc, after=40, align="left",
                  bottom_border=("B7791F", 14, 6))
        _run(p, subtitle, bold=True, color=GRAY, size=22)
        # 3. spacer
        p = _para(self.doc, after=260)
        _run(p, "\n", color=TEXT, size=20)
        # 4. big title
        p = _para(self.doc, after=60)
        _run(p, title, bold=True, color=NAVY, size=48)
        # 5. student name
        p = _para(self.doc, after=60)
        _run(p, student_name, bold=True, color=RED, size=34)
        # 6. cycle line
        p = _para(self.doc, after=300)
        _run(p, cycle_line, bold=False, color=GRAY, size=20)

    # -- profile info table (2-col, no header row) --------------------------
    def info_table(self, pairs):
        widths = [2700, 6660]
        tbl = self._new_table(len(pairs), widths)
        _tbl_borders(tbl)
        for i, (label, value) in enumerate(pairs):
            lc, vc = tbl.rows[i].cells
            _clear_cell(lc); _shade(lc, "E8ECF4"); _vcenter(lc)
            _cell_margins(lc, top=60, left=120, bottom=60, right=90)
            p = lc.add_paragraph()
            _sub(p._p.get_or_add_pPr(), "w:spacing", after=0, line=236)
            _run(p, label, bold=True, color=NAVY, size=18)
            fill = WHITE if i % 2 == 0 else "F4F4F4"
            _clear_cell(vc); _shade(vc, fill); _vcenter(vc)
            _cell_margins(vc, top=60, left=120, bottom=60, right=90)
            p = vc.add_paragraph()
            _sub(p._p.get_or_add_pPr(), "w:spacing", after=0, line=236)
            _run(p, value, bold=False, color=TEXT, size=18)
        return tbl

    def save(self, path):
        self.doc.save(path)
